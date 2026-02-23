#!/usr/bin/env python3
"""
pdf_to_docx.py — High-fidelity PDF to DOCX converter.

Extracts text (with font metadata), images, tables, equations, headers/footers
from each PDF and rebuilds as a Word document preserving layout and formatting.

Uses PyMuPDF (fitz) for PDF reading, python-docx for DOCX generation.
Output is saved to a 'docx_output/' subfolder.
"""

import fitz
import os
import re
import io
import sys
import shutil
from dataclasses import dataclass, field
from collections import Counter, defaultdict
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

# ─── Constants ─────────────────────────────────────────────────────────────────
HEADER_FOOTER_ZONE_RATIO = 0.08  # top/bottom 8% of page = header/footer zone
RUNNING_HEADER_RATIO = 0.40      # text on >40% of pages = running header (stricter)
SAMPLE_PAGES = 40                # pages to sample for header/footer detection
MAX_HEADER_FOOTER_LEN = 80       # headers/footers must be short text
MIN_IMAGE_DIM = 3              # min pixels to keep an image (filter 1x1 artifacts)
BLOCK_IMAGE_WIDTH_RATIO = 0.5  # image wider than 50% of text area = block image
TABLE_LINE_TOLERANCE = 3       # pt tolerance for line alignment
# Fonts ONLY used for math (never for body text)
MATH_ONLY_FONTS = {"cmmi", "cmsy", "cmex", "msbm", "msam", "eufm",
                   "rsfs", "bbm", "stmary", "wasy", "esint", "txsy",
                   "mtmi", "mtsyn", "mtex"}
# CM fonts used for BODY TEXT in TeX — NOT math indicators
# (CMR10 = roman text, CMBX10 = bold text, CMTI10 = italic text, etc.)
TEXT_CM_FONTS = {"cmr", "cmbx", "cmti", "cmss", "cmtt", "cmsl", "cmcsc"}
# Symbol/special-character fonts
SYMBOL_FONTS = {"symbol", "symbolmt", "wingdings", "zapfdingbats",
                "mathematica"}
# Combined set for is_math_font() — only truly math-indicating fonts
MATH_FONTS = MATH_ONLY_FONTS | SYMBOL_FONTS
MATH_CHARS = set(
    "∑∫∂√∞±≤≥≠≈∈∉⊂⊃⊆⊇∪∩∧∨¬∀∃∅∇∝∠∏⊗⊕⊥∥△▽"
    "αβγδεζηθικλμνξπρστυφχψω"
    "ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΠΡΣΤΥΦΧΨΩ"
    "₀₁₂₃₄₅₆₇₈₉⁰¹²³⁴⁵⁶⁷⁸⁹"
    "×÷·°′″⟨⟩≡≪≫ℏℓ∘≅∝∓∣∤≺≻≼≽"
    "⊢⊣⊤⊥⊨⊩⊲⊳⊴⊵⋀⋁⋂⋃⋄⋅⋆⋈⋮⋯⋰⋱"
    "←→↑↓↔↕↦⇐⇒⇔⇑⇓"
    "ℕℤℚℝℂℵ∂ℱℒℋ"
)
EQUATION_DPI = 200
FONT_SIZE_CLUSTER_THRESHOLD = 0.6
PARA_GAP_FACTOR = 1.8          # vertical gap > factor × font_size = new paragraph
BOLD_FONT_PATTERNS = ["bold", "cmbx", "advtib", "demi"]

# ─── Data Structures ──────────────────────────────────────────────────────────

@dataclass
class TextRun:
    """A span of text with uniform formatting."""
    text: str
    font_name: str
    font_size: float
    bold: bool
    italic: bool
    color: tuple = (0, 0, 0)
    superscript: bool = False
    subscript: bool = False
    is_math: bool = False


@dataclass
class PageElement:
    """Base for all elements on a page, sortable by position."""
    page_num: int
    bbox: tuple          # (x0, y0, x1, y1)
    element_type: str    # 'paragraph', 'image', 'table', 'equation'


@dataclass
class ParagraphElement(PageElement):
    runs: list = field(default_factory=list)
    alignment: str = "left"
    indent_left: float = 0
    indent_first: float = 0
    space_before: float = 0
    space_after: float = 0
    is_heading: bool = False
    heading_level: int = 0
    is_list_item: bool = False
    list_marker: str = ""

    def __post_init__(self):
        self.element_type = "paragraph"


@dataclass
class ImageElement(PageElement):
    image_data: bytes = b""
    image_ext: str = "png"
    width_pt: float = 0
    height_pt: float = 0
    is_block: bool = True

    def __post_init__(self):
        self.element_type = "image"


@dataclass
class TableCell:
    paragraphs: list = field(default_factory=list)  # list of ParagraphElement
    col_span: int = 1
    row_span: int = 1
    borders: dict = field(default_factory=lambda: {
        "top": True, "bottom": True, "left": True, "right": True
    })


@dataclass
class TableElement(PageElement):
    rows: list = field(default_factory=list)  # list of list of TableCell
    col_widths: list = field(default_factory=list)

    def __post_init__(self):
        self.element_type = "table"


@dataclass
class EquationElement(PageElement):
    image_data: bytes = b""
    width_pt: float = 0
    height_pt: float = 0
    is_display: bool = True  # display (block) vs inline

    def __post_init__(self):
        self.element_type = "equation"


# ─── Utility Functions ─────────────────────────────────────────────────────────

def is_bold_font(font_name: str, flags: int) -> bool:
    """Check if a font is bold by flags or name."""
    if flags & 16:  # bit 4 = bold
        return True
    name_lower = font_name.lower()
    return any(p in name_lower for p in BOLD_FONT_PATTERNS)


def is_italic_font(font_name: str, flags: int) -> bool:
    """Check if a font is italic by flags or name."""
    if flags & 2:  # bit 1 = italic
        return True
    name_lower = font_name.lower()
    return any(p in name_lower for p in ["italic", "oblique", "cmti", "slant"])


def is_math_font(font_name: str) -> bool:
    """Check if font is a math-indicating font (MATH_ONLY or SYMBOL).
    Does NOT return True for CM body text fonts (CMR10, CMBX10, CMTI10, etc.)
    which are used for regular text in TeX documents.
    """
    name_lower = font_name.lower()
    # Strip subset prefix like "ABCDEF+"
    if "+" in name_lower:
        name_lower = name_lower.split("+", 1)[1]
    return any(name_lower.startswith(mf) for mf in MATH_FONTS)


def has_math_chars(text: str) -> bool:
    """Check if text contains mathematical Unicode characters."""
    return bool(MATH_CHARS.intersection(text))


def bbox_overlap(a: tuple, b: tuple) -> bool:
    """Check if two bboxes overlap (x0, y0, x1, y1)."""
    return not (a[2] <= b[0] or a[0] >= b[2] or a[3] <= b[1] or a[1] >= b[3])


def bbox_contains(outer: tuple, inner: tuple, tolerance: float = 2) -> bool:
    """Check if outer bbox fully contains inner bbox."""
    return (inner[0] >= outer[0] - tolerance and
            inner[1] >= outer[1] - tolerance and
            inner[2] <= outer[2] + tolerance and
            inner[3] <= outer[3] + tolerance)


def pt_to_emu(pt: float) -> int:
    """Convert points to EMU."""
    return int(pt * 12700)


def color_from_value(val) -> tuple:
    """Extract RGB tuple (0-255) from fitz color value."""
    if val is None:
        return (0, 0, 0)
    if isinstance(val, (int, float)):
        c = int(val * 255)
        return (c, c, c)
    if isinstance(val, (list, tuple)):
        if len(val) == 1:
            c = int(val[0] * 255)
            return (c, c, c)
        if len(val) == 3:
            return tuple(int(c * 255) for c in val)
        if len(val) == 4:  # CMYK → approximate RGB
            c, m, y, k = val
            r = int(255 * (1 - c) * (1 - k))
            g = int(255 * (1 - m) * (1 - k))
            b = int(255 * (1 - y) * (1 - k))
            return (r, g, b)
    return (0, 0, 0)


def cluster_sizes(sizes: list, threshold: float = FONT_SIZE_CLUSTER_THRESHOLD) -> dict:
    """Cluster font sizes within threshold, return mapping size→canonical."""
    if not sizes:
        return {}
    sorted_sizes = sorted(set(sizes))
    clusters = []
    current = [sorted_sizes[0]]
    for s in sorted_sizes[1:]:
        if s - current[-1] <= threshold:
            current.append(s)
        else:
            clusters.append(current)
            current = [s]
    clusters.append(current)
    mapping = {}
    for cluster in clusters:
        canonical = sum(cluster) / len(cluster)
        for s in cluster:
            mapping[s] = canonical
    return mapping


def strip_subset_prefix(font_name: str) -> str:
    """Remove subset prefix like 'ABCDEF+' from font name."""
    if "+" in font_name and len(font_name.split("+")[0]) <= 7:
        return font_name.split("+", 1)[1]
    return font_name


def sanitize_text(text: str) -> str:
    """Remove characters that are invalid in XML 1.0 (used by python-docx).
    Keeps tabs, newlines, and all printable characters.
    """
    if not text:
        return text
    # XML 1.0 valid chars: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    cleaned = []
    for ch in text:
        cp = ord(ch)
        if cp == 0x9 or cp == 0xA or cp == 0xD:
            cleaned.append(ch)
        elif 0x20 <= cp <= 0xD7FF:
            cleaned.append(ch)
        elif 0xE000 <= cp <= 0xFFFD:
            cleaned.append(ch)
        elif 0x10000 <= cp <= 0x10FFFF:
            cleaned.append(ch)
        # else: skip invalid character
    return "".join(cleaned)


def is_scanned_pdf(doc, sample_size=10):
    """Detect if a PDF is mostly scanned images with no extractable text.
    Samples up to sample_size pages evenly distributed through the document.
    Returns True if <20% of sampled pages have meaningful text (>50 chars).
    """
    num_pages = len(doc)
    if num_pages == 0:
        return True
    if num_pages <= sample_size:
        indices = list(range(num_pages))
    else:
        indices = [int(i * num_pages / sample_size) for i in range(sample_size)]
    pages_with_text = 0
    for idx in indices:
        page = doc[idx]
        text = page.get_text("text").strip()
        if len(text) > 50:
            pages_with_text += 1
    ratio = pages_with_text / len(indices)
    return ratio < 0.2


# ─── Font Mapping ──────────────────────────────────────────────────────────────

FONT_MAP = {
    # PostScript standard fonts
    "times": "Times New Roman",
    "timesnewroman": "Times New Roman",
    "times-roman": "Times New Roman",
    "times-bold": "Times New Roman",
    "times-italic": "Times New Roman",
    "times-bolditalic": "Times New Roman",
    "helvetica": "Arial",
    "arial": "Arial",
    "arialmt": "Arial",
    "courier": "Courier New",
    "couriernew": "Courier New",
    "courier-bold": "Courier New",
    "verdana": "Verdana",
    "georgia": "Georgia",
    "palatino": "Palatino Linotype",
    "palatinolinotype": "Palatino Linotype",
    "bookman": "Bookman Old Style",
    "garamond": "Garamond",
    "cambria": "Cambria",
    "calibri": "Calibri",
    "tahoma": "Tahoma",
    "trebuchet": "Trebuchet MS",
    "trebuchetms": "Trebuchet MS",
    "symbol": "Symbol",
    "zapfdingbats": "Wingdings",
    # TeX/CM fonts → use serif fallback
    "cmr": "Times New Roman",
    "cmbx": "Times New Roman",
    "cmti": "Times New Roman",
    "cmss": "Arial",
    "cmtt": "Courier New",
    "cmmi": "Times New Roman",
    "cmsy": "Symbol",
    "cmex": "Symbol",
    # URW fonts (Ghostscript substitutions)
    "urwgothic": "Century Gothic",
    "urwbookman": "Bookman Old Style",
    "nimbus": "Times New Roman",
    "nimbussans": "Arial",
    "nimbusmono": "Courier New",
    "nimbusroman": "Times New Roman",
}


def map_font(pdf_font_name: str, flags: int = 0) -> str:
    """Map a PDF font name to the closest Word-compatible font."""
    clean = strip_subset_prefix(pdf_font_name).lower()
    # Remove style suffixes for lookup
    for suffix in ["-roman", "-bold", "-italic", "-bolditalic", "-oblique",
                   "-regular", "-light", "-medium", "-semibold", "-black",
                   "-demi", "-book", "mt", "-identity-h"]:
        clean = clean.replace(suffix, "")
    clean = clean.strip("-").strip()

    # Direct lookup
    if clean in FONT_MAP:
        return FONT_MAP[clean]

    # Prefix match
    for prefix, mapped in FONT_MAP.items():
        if clean.startswith(prefix):
            return mapped

    # Heuristic: serif vs sans-serif by flags
    # Flag bit 0 = fixed-pitch, bit 1 = serif, bit 2 = symbolic
    if flags & 2:  # serif
        return "Times New Roman"
    return "Arial"


# ─── Page Geometry Analysis ────────────────────────────────────────────────────

def analyze_page_geometry(doc: fitz.Document) -> tuple:
    """Analyze page sizes and margins from text bounding boxes.
    Returns (page_width_pt, page_height_pt, margins_dict).
    """
    page_widths = []
    page_heights = []
    text_lefts = []
    text_rights = []
    text_tops = []
    text_bottoms = []

    sample_count = min(20, len(doc))
    # Sample from throughout the document (skip first 2 pages which may be title/TOC)
    start = min(2, len(doc) - 1)
    step = max(1, (len(doc) - start) // sample_count)
    sample_indices = list(range(start, len(doc), step))[:sample_count]
    if not sample_indices:
        sample_indices = list(range(min(sample_count, len(doc))))

    for i in sample_indices:
        page = doc[i]
        rect = page.rect
        page_widths.append(rect.width)
        page_heights.append(rect.height)

        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] == 0:  # text block
                bb = block["bbox"]
                text_lefts.append(bb[0])
                text_rights.append(bb[2])
                text_tops.append(bb[1])
                text_bottoms.append(bb[3])

    if not text_lefts:
        # Fallback: use page dimensions with default margins
        pw = page_widths[0] if page_widths else 612
        ph = page_heights[0] if page_heights else 792
        return pw, ph, {"left": 72, "right": 72, "top": 72, "bottom": 72}

    # Use median page size
    page_widths.sort()
    page_heights.sort()
    pw = page_widths[len(page_widths) // 2]
    ph = page_heights[len(page_heights) // 2]

    # Margins: use 5th percentile for left/top (smallest observed text position)
    # and 95th percentile for right/bottom
    text_lefts.sort()
    text_rights.sort()
    text_tops.sort()
    text_bottoms.sort()

    idx_5 = max(0, len(text_lefts) // 20)
    idx_95 = min(len(text_rights) - 1, len(text_rights) * 19 // 20)

    left_margin = max(36, text_lefts[idx_5])
    right_margin = max(36, pw - text_rights[idx_95])
    top_margin = max(36, text_tops[idx_5])
    bottom_margin = max(36, ph - text_bottoms[idx_95])

    margins = {
        "left": left_margin,
        "right": right_margin,
        "top": top_margin,
        "bottom": bottom_margin,
    }
    return pw, ph, margins


# ─── Header/Footer Detection ──────────────────────────────────────────────────

def detect_headers_footers(doc: fitz.Document, body_size: float = 0) -> tuple:
    """Detect running headers and footers with strict criteria.

    A text block is only identified as a header/footer if ALL of:
    1. It appears on >40% of sampled pages (normalized form)
    2. It is in the top/bottom 8% of the page
    3. It is short (< 80 chars)
    4. If body_size is provided, its font size differs from body text

    Returns (header_patterns: set, footer_patterns: set) of normalized text.
    Does NOT return display text — we no longer recreate DOCX running headers.
    """
    sample_count = min(SAMPLE_PAGES, len(doc))
    start = min(3, len(doc) - 1)
    step = max(1, (len(doc) - start) // sample_count)
    indices = list(range(start, len(doc), step))[:sample_count]

    header_texts = Counter()
    footer_texts = Counter()

    for i in indices:
        page = doc[i]
        ph = page.rect.height
        header_zone = ph * HEADER_FOOTER_ZONE_RATIO
        footer_zone = ph * (1 - HEADER_FOOTER_ZONE_RATIO)
        blocks = page.get_text("dict")["blocks"]

        page_headers = set()
        page_footers = set()

        for block in blocks:
            if block["type"] != 0:
                continue
            bb = block["bbox"]
            text = ""
            block_font_size = 0
            span_count = 0
            for line in block["lines"]:
                for span in line["spans"]:
                    text += span["text"]
                    block_font_size += span.get("size", 0)
                    span_count += 1
            text = text.strip()

            # Must be short, non-empty text
            if not text or len(text) < 2 or len(text) > MAX_HEADER_FOOTER_LEN:
                continue

            # Check font size differs from body (if body_size known)
            if body_size > 0 and span_count > 0:
                avg_size = block_font_size / span_count
                if abs(avg_size - body_size) < 0.5:
                    continue  # Same size as body text — not a header

            normalized = re.sub(r"\d+", "N", text).strip()

            if bb[1] < header_zone:
                page_headers.add(normalized)
            elif bb[3] > footer_zone:
                page_footers.add(normalized)

        for t in page_headers:
            header_texts[t] += 1
        for t in page_footers:
            footer_texts[t] += 1

    threshold = max(3, int(sample_count * RUNNING_HEADER_RATIO))

    header_patterns = {t for t, c in header_texts.items() if c >= threshold}
    footer_patterns = {t for t, c in footer_texts.items() if c >= threshold}

    return header_patterns, footer_patterns


def is_in_header_footer_zone(bbox: tuple, page_height: float,
                              header_patterns: set, footer_patterns: set,
                              text: str) -> bool:
    """Check if text is in header/footer zone and matches a known pattern."""
    normalized = re.sub(r"\d+", "N", text).strip()
    header_zone = page_height * HEADER_FOOTER_ZONE_RATIO
    footer_zone = page_height * (1 - HEADER_FOOTER_ZONE_RATIO)
    if bbox[1] < header_zone and normalized in header_patterns:
        return True
    if bbox[3] > footer_zone and normalized in footer_patterns:
        return True
    return False


# ─── Image Extraction ──────────────────────────────────────────────────────────

def extract_page_images(page: fitz.Page, doc: fitz.Document,
                        text_area_width: float) -> list:
    """Extract images from a page with position info.
    Returns list of ImageElement.
    """
    elements = []
    page_num = page.number

    try:
        image_infos = page.get_image_info(xrefs=True)
    except Exception:
        return elements

    seen_xrefs = set()

    for info in image_infos:
        xref = info.get("xref", 0)
        if xref <= 0 or xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)

        # Get image bbox on page
        bbox = info.get("bbox", (0, 0, 0, 0))
        if not bbox or (bbox[2] - bbox[0] < MIN_IMAGE_DIM and
                        bbox[3] - bbox[1] < MIN_IMAGE_DIM):
            continue

        # Extract image data
        try:
            img_data = doc.extract_image(xref)
            if not img_data or not img_data.get("image"):
                continue
        except Exception:
            continue

        raw_bytes = img_data["image"]
        ext = img_data.get("ext", "png")
        if ext not in ("png", "jpeg", "jpg", "bmp", "gif", "tiff"):
            ext = "png"

        # Check actual image dimensions
        img_w = img_data.get("width", 0)
        img_h = img_data.get("height", 0)
        if img_w < MIN_IMAGE_DIM or img_h < MIN_IMAGE_DIM:
            continue

        # Width/height as rendered on page (in points)
        render_w = bbox[2] - bbox[0]
        render_h = bbox[3] - bbox[1]

        if render_w < 5 or render_h < 5:
            continue

        is_block = render_w > text_area_width * BLOCK_IMAGE_WIDTH_RATIO

        elements.append(ImageElement(
            page_num=page_num,
            bbox=bbox,
            element_type="image",
            image_data=raw_bytes,
            image_ext=ext,
            width_pt=render_w,
            height_pt=render_h,
            is_block=is_block,
        ))

    return elements


# ─── Table Detection ───────────────────────────────────────────────────────────

def detect_tables(page: fitz.Page) -> tuple:
    """Detect tables from vector line drawings.
    Returns (table_regions: list of bbox, table_elements: list of TableElement).
    """
    try:
        drawings = page.get_drawings()
    except Exception:
        return [], []

    if not drawings:
        return [], []

    h_lines = []  # (x0, y, x1) horizontal lines
    v_lines = []  # (x, y0, y1) vertical lines

    for d in drawings:
        for item in d.get("items", []):
            if item[0] == "l":  # line
                p1, p2 = item[1], item[2]
                x0, y0 = p1.x, p1.y
                x1, y1 = p2.x, p2.y

                # Horizontal line (y values close)
                if abs(y1 - y0) < TABLE_LINE_TOLERANCE and abs(x1 - x0) > 20:
                    h_lines.append((min(x0, x1), (y0 + y1) / 2, max(x0, x1)))
                # Vertical line (x values close)
                elif abs(x1 - x0) < TABLE_LINE_TOLERANCE and abs(y1 - y0) > 10:
                    v_lines.append(((x0 + x1) / 2, min(y0, y1), max(y0, y1)))

            elif item[0] == "re":  # rectangle — can form table borders
                rect = item[1]
                if isinstance(rect, fitz.Rect):
                    x0, y0, x1, y1 = rect.x0, rect.y0, rect.x1, rect.y1
                    w = x1 - x0
                    h = y1 - y0
                    # Thin rectangle → treat as lines
                    if h < 3 and w > 20:  # horizontal bar
                        h_lines.append((x0, (y0 + y1) / 2, x1))
                    elif w < 3 and h > 10:  # vertical bar
                        v_lines.append(((x0 + x1) / 2, y0, y1))
                    elif w > 20 and h > 10:  # full rect → 4 lines
                        h_lines.append((x0, y0, x1))
                        h_lines.append((x0, y1, x1))
                        v_lines.append((x0, y0, y1))
                        v_lines.append((x1, y0, y1))

    if len(h_lines) < 2 or len(v_lines) < 2:
        return [], []

    # Cluster y-coordinates of horizontal lines into rows
    h_ys = sorted(set(round(l[1], 0) for l in h_lines))
    v_xs = sorted(set(round(l[0], 0) for l in v_lines))

    # Merge close y-coordinates
    merged_ys = _merge_close_values(h_ys, TABLE_LINE_TOLERANCE * 2)
    merged_xs = _merge_close_values(v_xs, TABLE_LINE_TOLERANCE * 2)

    if len(merged_ys) < 2 or len(merged_xs) < 2:
        return [], []

    # Try to form table grids
    tables = []
    table_regions = []

    # Find contiguous grid regions
    grids = _find_table_grids(merged_xs, merged_ys, h_lines, v_lines)

    for grid_xs, grid_ys in grids:
        if len(grid_ys) < 2 or len(grid_xs) < 2:
            continue

        region = (grid_xs[0] - 2, grid_ys[0] - 2,
                  grid_xs[-1] + 2, grid_ys[-1] + 2)

        n_rows = len(grid_ys) - 1
        n_cols = len(grid_xs) - 1
        col_widths = [grid_xs[i + 1] - grid_xs[i] for i in range(n_cols)]

        rows = []
        for r in range(n_rows):
            row = []
            for c in range(n_cols):
                cell_bbox = (grid_xs[c], grid_ys[r],
                             grid_xs[c + 1], grid_ys[r + 1])
                cell = TableCell(
                    paragraphs=[],
                    borders={
                        "top": _has_line_at(h_lines, grid_ys[r],
                                            grid_xs[c], grid_xs[c + 1]),
                        "bottom": _has_line_at(h_lines, grid_ys[r + 1],
                                               grid_xs[c], grid_xs[c + 1]),
                        "left": _has_vline_at(v_lines, grid_xs[c],
                                              grid_ys[r], grid_ys[r + 1]),
                        "right": _has_vline_at(v_lines, grid_xs[c + 1],
                                               grid_ys[r], grid_ys[r + 1]),
                    }
                )
                row.append(cell)
            rows.append(row)

        table = TableElement(
            page_num=page.number,
            bbox=region,
            element_type="table",
            rows=rows,
            col_widths=col_widths,
        )
        tables.append(table)
        table_regions.append(region)

    return table_regions, tables


def _merge_close_values(values: list, tolerance: float) -> list:
    """Merge values that are within tolerance of each other."""
    if not values:
        return []
    merged = [values[0]]
    for v in values[1:]:
        if v - merged[-1] > tolerance:
            merged.append(v)
        else:
            merged[-1] = (merged[-1] + v) / 2
    return merged


def _find_table_grids(xs: list, ys: list, h_lines: list, v_lines: list) -> list:
    """Find rectangular grid regions from line intersections.
    Returns list of (grid_xs, grid_ys) tuples.
    """
    # Simplified: treat all detected lines as one potential grid
    # Validate that enough intersections exist
    valid_xs = []
    valid_ys = []

    for x in xs:
        # Check that this x-coordinate has vertical lines spanning multiple rows
        covering = sum(1 for vl in v_lines
                       if abs(vl[0] - x) < TABLE_LINE_TOLERANCE * 2)
        if covering >= 1:
            valid_xs.append(x)

    for y in ys:
        covering = sum(1 for hl in h_lines
                       if abs(hl[1] - y) < TABLE_LINE_TOLERANCE * 2)
        if covering >= 1:
            valid_ys.append(y)

    if len(valid_xs) >= 2 and len(valid_ys) >= 2:
        return [(valid_xs, valid_ys)]
    return []


def _has_line_at(h_lines: list, y: float, x0: float, x1: float) -> bool:
    """Check if a horizontal line exists at y spanning from x0 to x1."""
    tol = TABLE_LINE_TOLERANCE * 3
    for lx0, ly, lx1 in h_lines:
        if abs(ly - y) < tol and lx0 <= x0 + tol and lx1 >= x1 - tol:
            return True
    return False


def _has_vline_at(v_lines: list, x: float, y0: float, y1: float) -> bool:
    """Check if a vertical line exists at x spanning from y0 to y1."""
    tol = TABLE_LINE_TOLERANCE * 3
    for lx, ly0, ly1 in v_lines:
        if abs(lx - x) < tol and ly0 <= y0 + tol and ly1 >= y1 - tol:
            return True
    return False


def fill_table_cells(table: TableElement, page: fitz.Page,
                     header_patterns: set, footer_patterns: set) -> None:
    """Fill table cells with text extracted from the page."""
    grid_ys = []
    grid_xs = []
    # Reconstruct grid coordinates from col_widths and bbox
    x = table.bbox[0] + 2
    grid_xs.append(x)
    for w in table.col_widths:
        x += w
        grid_xs.append(x)

    y = table.bbox[1] + 2
    if table.rows:
        n_rows = len(table.rows)
        row_height = (table.bbox[3] - table.bbox[1]) / n_rows
        for r in range(n_rows + 1):
            grid_ys.append(table.bbox[1] + r * row_height)

    blocks = page.get_text("dict")["blocks"]
    for block in blocks:
        if block["type"] != 0:
            continue
        for line in block["lines"]:
            for span in line["spans"]:
                sp_bbox = span["bbox"]
                sp_cx = (sp_bbox[0] + sp_bbox[2]) / 2
                sp_cy = (sp_bbox[1] + sp_bbox[3]) / 2

                # Find which cell this span belongs to
                for r in range(len(table.rows)):
                    for c in range(len(table.rows[r])):
                        cell_x0 = grid_xs[c] if c < len(grid_xs) else table.bbox[0]
                        cell_x1 = grid_xs[c + 1] if c + 1 < len(grid_xs) else table.bbox[2]
                        cell_y0 = grid_ys[r] if r < len(grid_ys) else table.bbox[1]
                        cell_y1 = grid_ys[r + 1] if r + 1 < len(grid_ys) else table.bbox[3]

                        if cell_x0 <= sp_cx <= cell_x1 and cell_y0 <= sp_cy <= cell_y1:
                            text = span["text"].strip()
                            if text:
                                run = TextRun(
                                    text=text,
                                    font_name=span.get("font", ""),
                                    font_size=span.get("size", 10),
                                    bold=is_bold_font(span.get("font", ""),
                                                      span.get("flags", 0)),
                                    italic=is_italic_font(span.get("font", ""),
                                                          span.get("flags", 0)),
                                    color=color_from_value(span.get("color")),
                                )
                                # Add to cell
                                cell = table.rows[r][c]
                                if not cell.paragraphs:
                                    cell.paragraphs.append(ParagraphElement(
                                        page_num=page.number,
                                        bbox=sp_bbox,
                                        element_type="paragraph",
                                        runs=[],
                                    ))
                                cell.paragraphs[-1].runs.append(run)
                            break
                    else:
                        continue
                    break


# ─── Text Extraction & Paragraph Assembly ──────────────────────────────────────

def extract_text_elements(page: fitz.Page, table_regions: list,
                          header_patterns: set, footer_patterns: set,
                          body_size: float, heading_sizes: set) -> list:
    """Extract text from page, assemble into paragraphs.
    Excludes text in table regions and header/footer zones.
    Returns list of ParagraphElement.
    """
    page_height = page.rect.height
    page_width = page.rect.width
    blocks = page.get_text("dict")["blocks"]

    # Collect all text lines with metadata
    raw_lines = []  # list of dicts with keys: text, runs, bbox, y, baseline

    for block in blocks:
        if block["type"] != 0:
            continue
        block_bbox = block["bbox"]

        # Skip if block is inside a table region
        if any(bbox_overlap(block_bbox, tr) for tr in table_regions):
            continue

        # Pre-compute full block text once for header/footer check
        full_block_text = ""
        for l2 in block["lines"]:
            for s2 in l2["spans"]:
                full_block_text += s2["text"]

        # Check header/footer at block level (not per span)
        if is_in_header_footer_zone(block_bbox, page_height,
                                    header_patterns, footer_patterns,
                                    full_block_text.strip()):
            continue

        for line in block["lines"]:
            line_bbox = line["bbox"]
            line_text = ""
            line_runs = []

            for span in line["spans"]:
                text = span["text"]
                if not text:
                    continue

                font = span.get("font", "")
                size = span.get("size", 10)
                flags = span.get("flags", 0)
                color = color_from_value(span.get("color"))

                math = is_math_font(font) or has_math_chars(text)

                run = TextRun(
                    text=text,
                    font_name=font,
                    font_size=size,
                    bold=is_bold_font(font, flags),
                    italic=is_italic_font(font, flags),
                    color=color,
                    superscript=bool(flags & 1),
                    subscript=False,
                    is_math=math,
                )
                line_runs.append(run)
                line_text += text

            if line_runs and line_text.strip():
                raw_lines.append({
                    "text": line_text,
                    "runs": line_runs,
                    "bbox": line_bbox,
                    "y": line_bbox[1],
                    "x0": line_bbox[0],
                    "x1": line_bbox[2],
                    "font_size": line_runs[0].font_size if line_runs else 10,
                    "bold": any(r.bold for r in line_runs),
                    "has_math": any(r.is_math for r in line_runs),
                })

    if not raw_lines:
        return []

    # Sort by vertical position, then horizontal
    raw_lines.sort(key=lambda l: (l["y"], l["x0"]))

    # Assemble into paragraphs
    paragraphs = _assemble_paragraphs(raw_lines, page_width, body_size,
                                       heading_sizes)
    return paragraphs


def _assemble_paragraphs(lines: list, page_width: float,
                          body_size: float, heading_sizes: set) -> list:
    """Group lines into paragraphs using heuristics."""
    if not lines:
        return []

    paragraphs = []
    current_runs = list(lines[0]["runs"])
    current_bbox = list(lines[0]["bbox"])
    current_text = lines[0]["text"]
    current_size = lines[0]["font_size"]
    current_bold = lines[0]["bold"]

    for i in range(1, len(lines)):
        prev = lines[i - 1]
        curr = lines[i]

        should_join = _should_join_lines(prev, curr, body_size, heading_sizes,
                                          page_width)

        if should_join:
            # Join lines
            prev_text = prev["text"].rstrip()
            curr_text = curr["text"]

            # Handle hyphenation
            if prev_text.endswith("-") and curr_text and curr_text[0].islower():
                # Remove hyphen from last run
                if current_runs and current_runs[-1].text.endswith("-"):
                    current_runs[-1] = TextRun(
                        text=current_runs[-1].text[:-1],
                        font_name=current_runs[-1].font_name,
                        font_size=current_runs[-1].font_size,
                        bold=current_runs[-1].bold,
                        italic=current_runs[-1].italic,
                        color=current_runs[-1].color,
                        superscript=current_runs[-1].superscript,
                        subscript=current_runs[-1].subscript,
                        is_math=current_runs[-1].is_math,
                    )
            else:
                # Add space between joined lines if needed
                if (current_runs and not current_runs[-1].text.endswith(" ")
                        and curr["runs"] and not curr["runs"][0].text.startswith(" ")):
                    current_runs[-1] = TextRun(
                        text=current_runs[-1].text + " ",
                        font_name=current_runs[-1].font_name,
                        font_size=current_runs[-1].font_size,
                        bold=current_runs[-1].bold,
                        italic=current_runs[-1].italic,
                        color=current_runs[-1].color,
                        superscript=current_runs[-1].superscript,
                        subscript=current_runs[-1].subscript,
                        is_math=current_runs[-1].is_math,
                    )

            current_runs.extend(curr["runs"])
            current_text += " " + curr["text"]
            # Expand bbox
            current_bbox[2] = max(current_bbox[2], curr["bbox"][2])
            current_bbox[3] = curr["bbox"][3]
        else:
            # Finalize current paragraph
            para = _make_paragraph(current_runs, tuple(current_bbox),
                                    current_text, current_size, current_bold,
                                    body_size, heading_sizes, page_width,
                                    lines[i - 1]["bbox"])
            if para:
                paragraphs.append(para)

            # Start new paragraph
            current_runs = list(curr["runs"])
            current_bbox = list(curr["bbox"])
            current_text = curr["text"]
            current_size = curr["font_size"]
            current_bold = curr["bold"]

    # Don't forget the last paragraph
    if current_runs:
        para = _make_paragraph(current_runs, tuple(current_bbox),
                                current_text, current_size, current_bold,
                                body_size, heading_sizes, page_width,
                                lines[-1]["bbox"])
        if para:
            paragraphs.append(para)

    return paragraphs


def _should_join_lines(prev: dict, curr: dict, body_size: float,
                        heading_sizes: set, page_width: float) -> bool:
    """Decide if two consecutive text lines should be joined into one paragraph."""
    # Different font sizes → likely heading boundary
    if abs(prev["font_size"] - curr["font_size"]) > FONT_SIZE_CLUSTER_THRESHOLD:
        return False

    # Heading lines don't join with non-heading lines
    prev_is_heading = prev["font_size"] in heading_sizes or (
        prev["bold"] and prev["font_size"] > body_size + 1)
    curr_is_heading = curr["font_size"] in heading_sizes or (
        curr["bold"] and curr["font_size"] > body_size + 1)
    if prev_is_heading != curr_is_heading:
        return False

    # Vertical gap check
    gap = curr["y"] - prev["bbox"][3]
    line_height = prev["font_size"] * 1.2
    if gap > line_height * PARA_GAP_FACTOR:
        return False
    if gap < -line_height:  # overlapping or out of order
        return False

    # If both are headings with same font, they might be multi-line heading
    if prev_is_heading and curr_is_heading:
        if gap < line_height * 1.5:
            return True
        return False

    # Check for list item pattern
    list_pattern = re.compile(r"^(\d+[\.\)]\s|[•\-–■▪]\s|\([a-z]\)\s|\([iv]+\)\s)",
                              re.IGNORECASE)
    if list_pattern.match(curr["text"].strip()):
        return False

    # Left margin change → new paragraph or indented block
    margin_diff = abs(curr["x0"] - prev["x0"])
    if margin_diff > 20:  # significant indent change
        # But first line indent is ok if gap is small
        if gap > line_height * 0.3:
            return False

    # Previous line ends with sentence-ending punctuation AND next starts uppercase
    prev_text = prev["text"].rstrip()
    curr_text = curr["text"].strip()
    if (prev_text and prev_text[-1] in ".!?:" and
            curr_text and curr_text[0].isupper()):
        # Could be mid-sentence after abbreviation, or genuine sentence boundary
        # If gap is large, it's a new paragraph
        if gap > line_height * 0.5:
            return False

    # If prev line is short (< 60% of page width) and ends mid-line,
    # it might be end of paragraph (short last line)
    prev_width = prev["x1"] - prev["x0"]
    text_area = page_width * 0.7  # approximate text area width
    if prev_width < text_area * 0.6 and prev_text and prev_text[-1] in ".!?":
        return False

    # Default: join if gap is reasonable
    if gap <= line_height * 1.3:
        return True

    return False


def _make_paragraph(runs: list, bbox: tuple, text: str,
                     font_size: float, is_bold: bool,
                     body_size: float, heading_sizes: set,
                     page_width: float, last_line_bbox: tuple) -> Optional[ParagraphElement]:
    """Create a ParagraphElement from accumulated runs."""
    if not runs or not text.strip():
        return None

    # Detect heading
    is_heading = False
    heading_level = 0
    canonical_size = font_size
    if canonical_size in heading_sizes or (is_bold and canonical_size > body_size + 1):
        is_heading = True
        # Assign level based on size relative to heading sizes
        sorted_hsizes = sorted(heading_sizes, reverse=True)
        for idx, hs in enumerate(sorted_hsizes):
            if abs(canonical_size - hs) < FONT_SIZE_CLUSTER_THRESHOLD:
                heading_level = idx + 1
                break
        if heading_level == 0:
            heading_level = 1

    # Detect alignment
    center = page_width / 2
    text_center = (bbox[0] + bbox[2]) / 2
    text_width = bbox[2] - bbox[0]

    if text_width < page_width * 0.3 and abs(text_center - center) < 20:
        alignment = "center"
    elif text_width > page_width * 0.6:
        alignment = "justify"
    else:
        alignment = "left"

    # Detect list item
    is_list = False
    list_marker = ""
    stripped = text.strip()
    list_match = re.match(r"^(\d+[\.\)]\s|[•\-–■▪]\s|\([a-z]\)\s|\([iv]+\)\s)",
                          stripped, re.IGNORECASE)
    if list_match:
        is_list = True
        list_marker = list_match.group(1)

    # Indentation
    indent_left = max(0, bbox[0] - 72)  # relative to typical left margin

    page_num = 0  # will be set by caller

    return ParagraphElement(
        page_num=page_num,
        bbox=bbox,
        element_type="paragraph",
        runs=runs,
        alignment=alignment,
        indent_left=indent_left,
        indent_first=0,
        space_before=0,
        space_after=0,
        is_heading=is_heading,
        heading_level=heading_level,
        is_list_item=is_list,
        list_marker=list_marker,
    )


# ─── Equation Detection ───────────────────────────────────────────────────────

def detect_equations(page: fitz.Page, text_elements: list) -> list:
    """Detect math equations and render as images.
    Returns list of EquationElement (replacing text_elements that are equations).

    Optimization: merge consecutive math paragraphs into one image render
    to reduce the number of pixmap operations.
    """
    equations = []
    to_remove_indices = set()

    # First pass: identify which paragraphs are math
    math_indices = []
    for idx, para in enumerate(text_elements):
        if not isinstance(para, ParagraphElement):
            continue
        all_math = all(r.is_math for r in para.runs if r.text.strip())
        mostly_math = (sum(1 for r in para.runs if r.is_math) /
                       max(1, len(para.runs))) > 0.6
        if all_math or mostly_math:
            math_indices.append(idx)

    if not math_indices:
        return equations

    # Merge consecutive math paragraphs into groups for batch rendering
    groups = []
    current_group = [math_indices[0]]
    for i in range(1, len(math_indices)):
        prev_idx = math_indices[i - 1]
        curr_idx = math_indices[i]
        prev_bbox = text_elements[prev_idx].bbox
        curr_bbox = text_elements[curr_idx].bbox
        # If consecutive and close vertically (< 30pt gap), batch them
        if (curr_idx == prev_idx + 1 and
                curr_bbox[1] - prev_bbox[3] < 30):
            current_group.append(curr_idx)
        else:
            groups.append(current_group)
            current_group = [curr_idx]
    groups.append(current_group)

    # Render each group as a single image
    for group in groups:
        try:
            # Compute enclosing bbox for the group
            x0 = min(text_elements[i].bbox[0] for i in group)
            y0 = min(text_elements[i].bbox[1] for i in group)
            x1 = max(text_elements[i].bbox[2] for i in group)
            y1 = max(text_elements[i].bbox[3] for i in group)

            clip = fitz.Rect(x0 - 2, y0 - 2, x1 + 2, y1 + 2)
            clip = clip & page.rect

            pix = page.get_pixmap(clip=clip, dpi=EQUATION_DPI)
            img_bytes = pix.tobytes("png")

            width_pt = x1 - x0
            height_pt = y1 - y0

            eq = EquationElement(
                page_num=page.number,
                bbox=(x0, y0, x1, y1),
                element_type="equation",
                image_data=img_bytes,
                width_pt=width_pt,
                height_pt=height_pt,
                is_display=True,
            )
            equations.append(eq)
            to_remove_indices.update(group)
        except Exception:
            pass

    text_elements[:] = [e for i, e in enumerate(text_elements)
                        if i not in to_remove_indices]

    return equations


# ─── Vector Diagram Detection ─────────────────────────────────────────────────

def detect_diagrams(page: fitz.Page, table_regions: list) -> list:
    """Detect non-table vector drawing regions and render as images.
    Returns list of ImageElement for diagram regions.
    """
    try:
        drawings = page.get_drawings()
    except Exception:
        return []

    if not drawings:
        return []

    # Collect all drawing bboxes that are NOT in table regions
    drawing_bboxes = []
    for d in drawings:
        r = d.get("rect")
        if r is None:
            continue
        bbox = (r.x0, r.y0, r.x1, r.y1)
        # Skip tiny drawings
        if r.width < 20 or r.height < 20:
            continue
        # Skip if inside a table
        if any(bbox_overlap(bbox, tr) for tr in table_regions):
            continue
        drawing_bboxes.append(bbox)

    if not drawing_bboxes:
        return []

    # Cluster nearby drawings into diagram regions
    clusters = _cluster_bboxes(drawing_bboxes, gap_threshold=10)

    diagrams = []
    for cluster in clusters:
        # Compute enclosing bbox
        x0 = min(b[0] for b in cluster)
        y0 = min(b[1] for b in cluster)
        x1 = max(b[2] for b in cluster)
        y1 = max(b[3] for b in cluster)

        width = x1 - x0
        height = y1 - y0

        # Only render if the region is substantial (not just a line or separator)
        if width < 50 and height < 50:
            continue
        if width < 10 or height < 10:
            continue
        # Skip full-page width thin bars (separators)
        if height < 5:
            continue

        try:
            clip = fitz.Rect(x0 - 2, y0 - 2, x1 + 2, y1 + 2) & page.rect
            pix = page.get_pixmap(clip=clip, dpi=200)
            img_bytes = pix.tobytes("png")

            diagrams.append(ImageElement(
                page_num=page.number,
                bbox=(x0, y0, x1, y1),
                element_type="image",
                image_data=img_bytes,
                image_ext="png",
                width_pt=width,
                height_pt=height,
                is_block=True,
            ))
        except Exception:
            pass

    return diagrams


def _cluster_bboxes(bboxes: list, gap_threshold: float = 10) -> list:
    """Cluster overlapping or nearby bboxes into groups."""
    if not bboxes:
        return []

    # Simple union-find clustering
    n = len(bboxes)
    parent = list(range(n))

    def find(x):
        while parent[x] != x:
            parent[x] = parent[parent[x]]
            x = parent[x]
        return x

    def union(a, b):
        ra, rb = find(a), find(b)
        if ra != rb:
            parent[ra] = rb

    for i in range(n):
        for j in range(i + 1, n):
            # Check if bboxes overlap or are close
            a = bboxes[i]
            b = bboxes[j]
            expanded_a = (a[0] - gap_threshold, a[1] - gap_threshold,
                          a[2] + gap_threshold, a[3] + gap_threshold)
            if bbox_overlap(expanded_a, b):
                union(i, j)

    clusters = defaultdict(list)
    for i in range(n):
        clusters[find(i)].append(bboxes[i])

    return list(clusters.values())


# ─── Body Size / Heading Size Detection ────────────────────────────────────────

def analyze_document_fonts(doc: fitz.Document) -> tuple:
    """Analyze the document to find body text size and heading sizes.
    Returns (body_size, heading_sizes_set).
    """
    size_counter = Counter()
    bold_sizes = set()

    sample = min(50, len(doc))
    step = max(1, len(doc) // sample)
    for i in range(0, len(doc), step):
        if i >= len(doc):
            break
        page = doc[i]
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    if len(text) < 3:
                        continue
                    size = round(span["size"], 1)
                    if size < 4:
                        continue
                    size_counter[size] += len(text)
                    if is_bold_font(span.get("font", ""), span.get("flags", 0)):
                        bold_sizes.add(size)

    if not size_counter:
        return 10, set()

    # Body size = most common size by character count
    body_size = size_counter.most_common(1)[0][0]

    # Heading sizes = bold sizes larger than body
    size_map = cluster_sizes(list(size_counter.keys()))
    canonical_body = size_map.get(body_size, body_size)

    heading_sizes = set()
    for s in bold_sizes:
        canonical = size_map.get(s, s)
        if canonical > canonical_body + 0.5:
            heading_sizes.add(s)

    return body_size, heading_sizes


# ─── DOCX Builder ──────────────────────────────────────────────────────────────

def setup_page_layout(doc_docx: Document, page_width: float,
                       page_height: float, margins: dict) -> None:
    """Configure DOCX page size and margins."""
    section = doc_docx.sections[0]
    section.page_width = pt_to_emu(page_width)
    section.page_height = pt_to_emu(page_height)
    section.left_margin = pt_to_emu(margins["left"])
    section.right_margin = pt_to_emu(margins["right"])
    section.top_margin = pt_to_emu(margins["top"])
    section.bottom_margin = pt_to_emu(margins["bottom"])


def render_paragraph(doc_docx: Document, para_elem: ParagraphElement,
                      left_margin: float) -> None:
    """Render a paragraph element to DOCX."""
    p = doc_docx.add_paragraph()

    # Alignment
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    p.alignment = align_map.get(para_elem.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # Paragraph formatting
    pf = p.paragraph_format
    if para_elem.indent_left > 10:
        pf.left_indent = Pt(min(para_elem.indent_left - left_margin, 200))
    if para_elem.space_before > 0:
        pf.space_before = Pt(para_elem.space_before)
    if para_elem.space_after > 0:
        pf.space_after = Pt(para_elem.space_after)

    # Heading style
    if para_elem.is_heading and 1 <= para_elem.heading_level <= 9:
        try:
            p.style = doc_docx.styles[f"Heading {para_elem.heading_level}"]
        except (KeyError, ValueError):
            pass

    # Add runs
    for run_data in para_elem.runs:
        r = p.add_run(sanitize_text(run_data.text))
        r.font.name = map_font(run_data.font_name)
        r.font.size = Pt(run_data.font_size)
        r.font.bold = run_data.bold
        r.font.italic = run_data.italic
        if run_data.color != (0, 0, 0):
            try:
                r.font.color.rgb = RGBColor(*run_data.color)
            except (ValueError, TypeError):
                pass
        if run_data.superscript:
            r.font.superscript = True
        if run_data.subscript:
            r.font.subscript = True

    return p


def render_image(doc_docx: Document, img_elem) -> None:
    """Render an image element to DOCX."""
    try:
        stream = io.BytesIO(img_elem.image_data)
        # Validate image
        try:
            pil_img = PILImage.open(io.BytesIO(img_elem.image_data))
            pil_img.verify()
        except Exception:
            return

        p = doc_docx.add_paragraph()
        if img_elem.is_block if hasattr(img_elem, 'is_block') else True:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        # Calculate size — cap at page width
        section = doc_docx.sections[-1]
        max_width_emu = (section.page_width - section.left_margin -
                          section.right_margin)
        max_width_pt = max_width_emu / 12700

        width = min(img_elem.width_pt, max_width_pt)
        scale = width / img_elem.width_pt if img_elem.width_pt > 0 else 1
        height = img_elem.height_pt * scale

        run.add_picture(stream, width=Pt(width), height=Pt(height))
    except Exception as e:
        # If image embedding fails, skip silently
        pass


def render_table(doc_docx: Document, table_elem: TableElement) -> None:
    """Render a table element to DOCX."""
    if not table_elem.rows:
        return

    n_rows = len(table_elem.rows)
    n_cols = max(len(row) for row in table_elem.rows) if table_elem.rows else 0

    if n_rows == 0 or n_cols == 0:
        return

    table = doc_docx.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    section = doc_docx.sections[-1]
    available_width = (section.page_width - section.left_margin -
                        section.right_margin)

    if table_elem.col_widths:
        total_pdf_width = sum(table_elem.col_widths)
        if total_pdf_width > 0:
            for i, w in enumerate(table_elem.col_widths):
                if i < n_cols:
                    ratio = w / total_pdf_width
                    col_emu = int(available_width * ratio)
                    for row in table.rows:
                        if i < len(row.cells):
                            row.cells[i].width = col_emu

    # Fill cells
    for r, row_data in enumerate(table_elem.rows):
        for c, cell_data in enumerate(row_data):
            if r >= n_rows or c >= n_cols:
                continue
            cell = table.rows[r].cells[c]

            # Clear default paragraph
            if cell.paragraphs:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()

            for para in cell_data.paragraphs:
                for run_data in para.runs:
                    run = p.add_run(sanitize_text(run_data.text) + " ")
                    run.font.name = map_font(run_data.font_name)
                    run.font.size = Pt(min(run_data.font_size, 11))
                    run.font.bold = run_data.bold
                    run.font.italic = run_data.italic

            # Set borders via XML
            _set_cell_borders(cell, cell_data.borders)


def _set_cell_borders(cell, borders: dict) -> None:
    """Set cell borders via XML manipulation."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing borders
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)

    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        if borders.get(edge, True):
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
        tcBorders.append(el)

    tcPr.append(tcBorders)


def render_equation(doc_docx: Document, eq_elem: EquationElement) -> None:
    """Render an equation as an image."""
    render_image(doc_docx, ImageElement(
        page_num=eq_elem.page_num,
        bbox=eq_elem.bbox,
        element_type="image",
        image_data=eq_elem.image_data,
        image_ext="png",
        width_pt=eq_elem.width_pt,
        height_pt=eq_elem.height_pt,
        is_block=eq_elem.is_display,
    ))


def render_element(doc_docx: Document, element: PageElement,
                    left_margin: float) -> None:
    """Render any page element to DOCX."""
    if isinstance(element, ParagraphElement):
        render_paragraph(doc_docx, element, left_margin)
    elif isinstance(element, ImageElement):
        render_image(doc_docx, element)
    elif isinstance(element, TableElement):
        render_table(doc_docx, element)
    elif isinstance(element, EquationElement):
        render_equation(doc_docx, element)


# ─── Main Processing ──────────────────────────────────────────────────────────

def sort_elements(elements: list) -> list:
    """Sort elements by vertical position (top of bbox), then horizontal."""
    return sorted(elements, key=lambda e: (e.bbox[1], e.bbox[0]))


def process_pdf(pdf_path: str, output_path: str) -> str:
    """Convert a single PDF to DOCX.
    Returns 'success', 'scanned', or 'failed'.
    """
    filename = os.path.basename(pdf_path)
    print(f"\n{'='*70}")
    print(f"Processing: {filename}")
    print(f"{'='*70}")

    try:
        doc_pdf = fitz.open(pdf_path)
    except Exception as e:
        print(f"  ERROR: Cannot open PDF: {e}")
        return "failed"

    num_pages = len(doc_pdf)
    print(f"  Pages: {num_pages}")

    # Check for scanned/image-only PDF
    if is_scanned_pdf(doc_pdf):
        print(f"  SKIPPED: scanned/image-only PDF (no extractable text)")
        doc_pdf.close()
        return "scanned"

    # Step 1: Analyze page geometry
    print("  [1/7] Analyzing page geometry...")
    page_width, page_height, margins = analyze_page_geometry(doc_pdf)
    print(f"         Page: {page_width:.0f}×{page_height:.0f}pt, "
          f"margins: L={margins['left']:.0f} R={margins['right']:.0f} "
          f"T={margins['top']:.0f} B={margins['bottom']:.0f}")

    text_area_width = page_width - margins["left"] - margins["right"]

    # Step 2: Analyze fonts (needed before header detection for font-size filter)
    print("  [2/7] Analyzing document fonts...")
    body_size, heading_sizes = analyze_document_fonts(doc_pdf)
    print(f"         Body: {body_size:.1f}pt, "
          f"Heading sizes: {sorted(heading_sizes)[:5]}")

    # Step 3: Detect headers/footers (with body_size for font-size filtering)
    print("  [3/7] Detecting headers/footers...")
    header_patterns, footer_patterns = detect_headers_footers(doc_pdf, body_size)
    print(f"         Headers: {len(header_patterns)} patterns, "
          f"Footers: {len(footer_patterns)} patterns")

    # Step 4: Create DOCX
    doc_docx = Document()
    setup_page_layout(doc_docx, page_width, page_height, margins)

    # Remove the default empty paragraph
    if doc_docx.paragraphs:
        p = doc_docx.paragraphs[0]
        p_element = p._element
        p_element.getparent().remove(p_element)

    # Step 5: Process each page
    print(f"  [4/7] Processing {num_pages} pages...")
    total_elements = 0

    for page_num in range(num_pages):
        page = doc_pdf[page_num]

        if page_num > 0 and page_num % 50 == 0:
            print(f"         Page {page_num}/{num_pages}...")

        try:
            # Page break between pages (except first)
            if page_num > 0:
                p = doc_docx.add_paragraph()
                p.paragraph_format.page_break_before = True

            # Extract images
            images = extract_page_images(page, doc_pdf, text_area_width)

            # Detect tables
            table_regions, tables = detect_tables(page)
            for table in tables:
                fill_table_cells(table, page, header_patterns, footer_patterns)

            # Extract text (excluding tables and headers/footers)
            text_elements = extract_text_elements(
                page, table_regions, header_patterns, footer_patterns,
                body_size, heading_sizes
            )

            # Set page_num on text elements
            for elem in text_elements:
                elem.page_num = page_num

            # Detect equations within text
            equations = detect_equations(page, text_elements)

            # Combine and sort all elements
            all_elements = text_elements + images + tables + equations
            all_elements = sort_elements(all_elements)

            # Deduplicate overlapping elements
            all_elements = _deduplicate_elements(all_elements)

            total_elements += len(all_elements)

            # Render
            for element in all_elements:
                render_element(doc_docx, element, margins["left"])

        except Exception as e:
            print(f"         WARNING: Page {page_num} error: {e}")
            # Add a placeholder paragraph so page flow isn't disrupted
            try:
                p = doc_docx.add_paragraph()
                p.add_run(f"[Page {page_num + 1} - conversion error]")
            except Exception:
                pass

    print(f"         Total elements rendered: {total_elements}")

    # Step 6: Save
    print(f"  [5/7] Saving DOCX...")
    try:
        doc_docx.save(output_path)
        size_mb = os.path.getsize(output_path) / (1024 * 1024)
        print(f"  [6/7] Saved: {output_path} ({size_mb:.1f} MB)")
    except Exception as e:
        print(f"  ERROR saving: {e}")
        return "failed"

    doc_pdf.close()
    print(f"  [7/7] Done!")
    return "success"


def _deduplicate_elements(elements: list) -> list:
    """Remove duplicate/overlapping elements, preferring text over images."""
    if len(elements) <= 1:
        return elements

    result = []
    used_regions = []

    for elem in elements:
        # Check if this element's region is already covered
        overlaps = False
        for used_bbox, used_type in used_regions:
            if bbox_overlap(elem.bbox, used_bbox):
                # If both are images or both are text, skip the duplicate
                if elem.element_type == used_type:
                    overlaps = True
                    break
                # If one is equation/diagram image and other is text paragraph,
                # the one already placed takes precedence
                if elem.element_type == "image" and used_type in ("paragraph", "equation"):
                    overlaps = True
                    break

        if not overlaps:
            result.append(elem)
            used_regions.append((elem.bbox, elem.element_type))

    return result


def main():
    """Process all PDFs in the working directory."""
    # Ensure unbuffered output
    sys.stdout.reconfigure(line_buffering=True)

    work_dir = Path("/Users/siraj/Downloads/For_TOC_Make")
    output_dir = work_dir / "docx_output"
    output_dir.mkdir(exist_ok=True)

    pdf_files = sorted(work_dir.glob("*.pdf"))
    if not pdf_files:
        print("No PDF files found in", work_dir)
        sys.exit(1)

    failed_dir = work_dir / "failed_files"

    print(f"Found {len(pdf_files)} PDF files")
    print(f"Output directory: {output_dir}")

    results = {"success": [], "failed": [], "scanned": []}

    for pdf_path in pdf_files:
        docx_name = pdf_path.stem + ".docx"
        output_path = output_dir / docx_name

        try:
            result = process_pdf(str(pdf_path), str(output_path))
            if result == "success":
                results["success"].append(pdf_path.name)
            elif result == "scanned":
                results["scanned"].append(pdf_path.name)
                failed_dir.mkdir(exist_ok=True)
                shutil.copy2(str(pdf_path), str(failed_dir / pdf_path.name))
                print(f"  → Copied to failed_files/{pdf_path.name}")
            else:
                results["failed"].append(pdf_path.name)
        except Exception as e:
            import traceback
            print(f"  EXCEPTION: {e}")
            traceback.print_exc()
            results["failed"].append(pdf_path.name)

    # Summary
    total = len(pdf_files)
    print(f"\n{'='*70}")
    print("SUMMARY")
    print(f"{'='*70}")
    print(f"  Succeeded: {len(results['success'])}/{total}")
    for name in results["success"]:
        print(f"    ✓ {name}")
    if results["scanned"]:
        print(f"  Scanned (→ failed_files/): {len(results['scanned'])}/{total}")
        for name in results["scanned"]:
            print(f"    ⚠ {name}")
    if results["failed"]:
        print(f"  Failed: {len(results['failed'])}/{total}")
        for name in results["failed"]:
            print(f"    ✗ {name}")


if __name__ == "__main__":
    main()
